"""Multi-format text extraction for API documents.

Supports: plain text, Markdown, YAML, JSON, PDF, DOCX, DOC, HTML.
Provides a single ``extract_text()`` entry point.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """Extract plain text from a document file of any supported format.

    Returns the full text content as a string. Returns empty string on failure
    (caller should check and raise if appropriate).
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    try:
        if ext in (".txt", ".md", ".markdown", ".yaml", ".yml", ".json"):
            return path.read_text(encoding="utf-8")

        if ext == ".pdf":
            from doc_parser.pdf_parser import PdfParser
            return PdfParser.parse(file_path)

        if ext == ".docx":
            return _extract_docx(file_path)

        if ext == ".doc":
            return _extract_doc(file_path)

        if ext in (".html", ".htm"):
            return _extract_html(file_path)

        # Unknown extension — try as plain text first
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raise ValueError(
                f"Unsupported binary format: {file_path}. "
                f"Only text-based formats are supported (txt, md, yaml, json, "
                f"pdf, docx, doc, html). Binary files (images, etc.) cannot be processed."
            ) from None
    except ValueError:
        raise
    except Exception as e:
        logger.warning("Failed to extract text from %s: %s", file_path, e)
        return ""


# ------------------------------------------------------------------
# Format-specific helpers
# ------------------------------------------------------------------

def _extract_docx(file_path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install: pip install python-docx"
        )

    doc = Document(file_path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables (common for API docs)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _extract_doc(file_path: str) -> str:
    """Extract text from a legacy .doc file.

    Tries antiword first (if on PATH), then falls back to treating
    it as a .docx file (some .doc files are actually .docx).
    """
    import shutil
    import subprocess

    # Try antiword
    if shutil.which("antiword"):
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except Exception:
            pass

    # Try as DOCX (some .doc files are actually DOCX format)
    try:
        return _extract_docx(file_path)
    except Exception:
        pass

    raise RuntimeError(
        "Cannot parse legacy .doc files. "
        "Install antiword (apt install antiword) or convert to .docx."
    )


def _extract_html(file_path: str) -> str:
    """Extract text from an HTML file.

    Uses BeautifulSoup if available, otherwise simple regex tag stripping.
    """
    content = Path(file_path).read_text(encoding="utf-8")

    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        # Fallback: regex-based tag stripping
        content = re.sub(
            r'<script[^>]*>.*?</script>', '', content,
            flags=re.DOTALL | re.IGNORECASE,
        )
        content = re.sub(
            r'<style[^>]*>.*?</style>', '', content,
            flags=re.DOTALL | re.IGNORECASE,
        )
        content = re.sub(r'<[^>]+>', ' ', content)
        content = re.sub(r'\s+', ' ', content)
        return content.strip()
